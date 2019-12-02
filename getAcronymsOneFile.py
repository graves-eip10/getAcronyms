import re
from docx import Document
from openpyxl import Workbook

print("Please enter the document name:")
docName = input()

doc = Document(docName)

def getCapsInParentheses(myDoc):
	#Prepare an accumulator for the file contents
	fullText = ""

	#Loop over the document paragraph and table objects, adding the text to the accumulator fullText.
	for paragraph in myDoc.paragraphs:
		fullText += paragraph.text

	for table in myDoc.tables:
		for row in table.rows:
			for cell in row.cells:
				fullText += cell.text

	#Parse the full text of the document and return only the caps in parentheses
	capsInParentheses = re.findall('\(([A-Z]+)\)', fullText)
	
	return capsInParentheses

capsInParentheses = getCapsInParentheses(doc)

#Prepare an Excel Workbook to store a table of the arconyms found.
acronymSink = Workbook()
ws = acronymSink.active
ws.title = "Acronyms"

ws.cell(row = 1, column = 1).value = 'Acronyms'

#Prepare a counter for the excel rows
rowIndex = 2

for i in range(0,len(capsInParentheses)):
	ws.cell(row = rowIndex, column = 1).value = capsInParentheses[i]
	rowIndex += 1

acronymSink.save('output.xlsx')
print("Possible acronyms written to file.")